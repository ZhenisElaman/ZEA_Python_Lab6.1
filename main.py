from docx import Document
from docx.shared import Inches
guestnames = ['Zbrush', 'Maya', 'Blender', '3DMax']
itr = 0
for x in guestnames:
 document = Document()

 document.add_picture('logo.jpg', width=Inches(0.5))
 document.add_heading('3D Character artist - Zhenis Elaman', 0)

 if itr == 0: n = 5
 elif itr == 1: n = 6
 elif itr == 2: n = 7
 elif itr == 3: n = 8

 p = document.add_paragraph(str(n) + ' лет опыта работы в ' + x)
 p.add_run(' Активно ищу работу').bold = True
 p.add_run(', на удаленной основе или с релокацией.')
 document.add_heading('Я жду ваши заявки!!!', level = 1)
 document.add_paragraph('')
 document.add_picture('skull.jpg', width=Inches(5.25))
 document.add_page_break()
 document.save(str(itr) + '-' + x + '.docx')
 itr+=1